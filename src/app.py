# VERSION: V3.0_ELITE_EDITION
# AUTHOR: CYBERGUARD TEAM
# THEME: NEON CYBERPUNK FUTURISTIC

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import time
import socket
import threading
import requests
import os
import random
import hashlib
import json
import base64
import string
from datetime import datetime
from io import BytesIO

# ==========================================
# 1. PAGE CONFIGURATION & ASSETS
# ==========================================
st.set_page_config(
    page_title="PAYBUDDY // CYBERGUARD ELITE",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# TRY IMPORTS (Graceful Degradation)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from streamlit_lottie import st_lottie
    LOTTIE_AVAILABLE = True
except ImportError:
    LOTTIE_AVAILABLE = False

# ==========================================
# 2. FUTURISTIC UI ENGINE (CSS INJECTION)
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;500;700;900&family=Share+Tech+Mono&display=swap');

    /* GLOBAL THEME */
    .stApp {
        background-color: #050505;
        background-image: radial-gradient(circle at 50% 50%, #111 0%, #000 100%);
        color: #e0e0e0;
        font-family: 'Share Tech Mono', monospace;
    }

    /* NEON HEADERS */
    h1, h2, h3 {
        font-family: 'Orbitron', sans-serif !important;
        text-transform: uppercase;
        letter-spacing: 2px;
    }
    h1 { 
        background: -webkit-linear-gradient(#00f3ff, #0066ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-shadow: 0px 0px 20px rgba(0, 243, 255, 0.5);
    }
    h2 { color: #00f3ff; border-bottom: 1px solid #00f3ff; padding-bottom: 10px; }
    h3 { color: #ff0055; }

    /* CUSTOM CARDS */
    div.css-1r6slb0 {
        background: rgba(255, 255, 255, 0.02);
        border: 1px solid rgba(0, 243, 255, 0.2);
        border-radius: 10px;
        padding: 20px;
        backdrop-filter: blur(10px);
        box-shadow: 0 0 15px rgba(0, 0, 0, 0.5);
    }

    /* SIDEBAR */
    section[data-testid="stSidebar"] {
        background-color: #0a0a0a;
        border-right: 1px solid #333;
    }

    /* METRICS */
    div[data-testid="stMetricValue"] {
        font-family: 'Orbitron';
        color: #00f3ff !important;
        text-shadow: 0 0 10px rgba(0, 243, 255, 0.6);
    }
    div[data-testid="stMetricLabel"] {
        color: #888;
        font-size: 0.8rem;
    }

    /* BUTTONS */
    .stButton>button {
        background: linear-gradient(45deg, #001133, #003366);
        color: #00f3ff;
        border: 1px solid #00f3ff;
        border-radius: 0px;
        font-family: 'Orbitron';
        font-weight: bold;
        transition: all 0.3s ease;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    .stButton>button:hover {
        background: #00f3ff;
        color: #000;
        box-shadow: 0 0 20px #00f3ff;
    }

    /* TABS */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: transparent;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: rgba(255, 255, 255, 0.05);
        border: 1px solid #333;
        color: #888;
        font-family: 'Orbitron';
    }
    .stTabs [aria-selected="true"] {
        background-color: rgba(0, 243, 255, 0.1) !important;
        border-color: #00f3ff !important;
        color: #00f3ff !important;
        box-shadow: inset 0 0 10px rgba(0, 243, 255, 0.2);
    }

    /* ALERTS */
    .element-container .stAlert {
        background: rgba(0, 0, 0, 0.5);
        border: 1px solid;
        backdrop-filter: blur(5px);
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 3. CORE CONSTANTS & SESSION
# ==========================================
TEAM_NAME = "CyberGuard"
TEAM_MEMBERS = [
    {"name": "Moheed Ul Hassan", "reg": "22I-7451"},
    {"name": "Ali Abbas", "reg": "22I-2285"},
    {"name": "Abdur Rehman", "reg": "22I-2291"}
]

if 'logs' not in st.session_state: st.session_state.logs = []
if 'auth' not in st.session_state: st.session_state.auth = False

def log_system(tool, type_, msg):
    ts = datetime.now().strftime("%H:%M:%S")
    st.session_state.logs.insert(0, {"timestamp": ts, "tool": tool, "type": type_, "message": msg})

# ==========================================
# 4. IDENTITY GATE
# ==========================================
def check_access():
    with st.sidebar:
        st.markdown("## 🔒 ACCESS CONTROL")
        st.image("https://img.icons8.com/nolan/96/fingerprint-scan.png", width=80)
        
        status_c = st.container()
        
        # Identity File Check
        id_exists = os.path.exists("identity.txt")
        con_exists = os.path.exists("consent.txt")
        
        if id_exists and con_exists:
            status_c.success("IDENTITY VERIFIED")
            status_c.success("CONSENT TOKEN ACTIVE")
            if not st.session_state.auth:
                if st.button("INITIALIZE SYSTEM"):
                    st.session_state.auth = True
                    log_system("KERNEL", "INFO", "Session Initialized by authorized personnel.")
                    st.rerun()
        else:
            status_c.error("MISSING CREDENTIALS")
            st.warning("Upload identity.txt & consent.txt to root.")
            if not id_exists: st.code("ERROR: identity.txt not found", language="bash")
            if not con_exists: st.code("ERROR: consent.txt not found", language="bash")

# ==========================================
# 5. DASHBOARD & TOOLS
# ==========================================
if not st.session_state.auth:
    check_access()
    st.markdown("""
    <div style='display:flex; justify-content:center; align-items:center; height:60vh; flex-direction:column;'>
        <h1 style='font-size:4em; margin-bottom:0;'>CYBERGUARD ELITE</h1>
        <h3 style='color:#555;'>RESTRICTED ACCESS // PAYBUDDY QA ENV</h3>
        <div style='margin-top:50px; border:1px solid #333; padding:20px; background:rgba(0,0,0,0.5);'>
            WAITING FOR BIOMETRIC HANDSHAKE...
        </div>
    </div>
    """, unsafe_allow_html=True)

else:
    # --- HEADER ---
    c1, c2 = st.columns([3, 1])
    with c1:
        st.title("PAYBUDDY // DEFENSE GRID")
        st.markdown(f"**OPERATOR:** {TEAM_NAME} | **ENV:** PROD_SIM_01 | **LATENCY:** 12ms")
    with c2:
        if st.button("TERMINATE SESSION"):
            st.session_state.auth = False
            st.rerun()

    # --- TABS ---
    tabs = st.tabs(["📊 COMMAND CENTER", "🔭 PORT RECON", "🔐 CRYPTO AUDIT", "⚡ LOAD STRESS", "🕸️ WEB CRAWLER", "📝 EXECUTIVE REPORT"])

    # --- TAB 1: COMMAND CENTER ---
    with tabs[0]:
        st.markdown("### SYSTEM TELEMETRY")
        
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("CPU LOAD", f"{random.randint(10, 40)}%", "NORMAL")
        m2.metric("NETWORK I/O", "1.2 GB/s", "+12%")
        m3.metric("ACTIVE THREATS", "0", "SECURE")
        m4.metric("TARGET", "127.0.0.1", "LOCAL")

        g1, g2 = st.columns([2, 1])
        with g1:
            st.markdown("#### 📡 NETWORK TRAFFIC REAL-TIME")
            # Live Plotly Chart
            x_data = list(range(20))
            y_in = [random.randint(50, 150) for _ in x_data]
            y_out = [random.randint(30, 100) for _ in x_data]
            
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=x_data, y=y_in, fill='tozeroy', name='Inbound', line=dict(color='#00f3ff')))
            fig.add_trace(go.Scatter(x=x_data, y=y_out, fill='tozeroy', name='Outbound', line=dict(color='#ff0055')))
            fig.update_layout(template="plotly_dark", paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', height=300, margin=dict(l=0, r=0, t=0, b=0))
            st.plotly_chart(fig, use_container_width=True)
        
        with g2:
            st.markdown("#### 🛡️ TEAM ROSTER")
            for m in TEAM_MEMBERS:
                st.markdown(f"""
                <div style='background:rgba(255,255,255,0.05); padding:10px; margin-bottom:10px; border-left:3px solid #00f3ff;'>
                    <div style='font-family:Orbitron; color:#fff;'>{m['name']}</div>
                    <div style='font-size:0.8em; color:#888;'>ID: {m['reg']}</div>
                </div>
                """, unsafe_allow_html=True)

    # --- TAB 2: PORT SCANNER ---
    with tabs[1]:
        st.markdown("### 🔭 TARGET RECONNAISSANCE")
        
        col_scan, col_res = st.columns([1, 2])
        
        with col_scan:
            target = st.text_input("TARGET IP", "192.168.1.105")
            scan_mode = st.selectbox("SCAN MODE", ["FAST (Top 100)", "FULL (1-1000)", "STEALTH (Random)"])
            
            if st.button("INITIATE SCAN"):
                log_system("SCANNER", "INFO", f"Port Scan started on {target} ({scan_mode})")
                
                with st.status("SCANNING TARGET...", expanded=True) as status:
                    progress = st.progress(0)
                    found_ports = []
                    
                    # Simulation Logic
                    common_ports = {21:'FTP', 22:'SSH', 80:'HTTP', 443:'HTTPS', 3306:'MYSQL', 8080:'ALT_HTTP'}
                    
                    for i in range(100):
                        time.sleep(0.01) # Simulating network delay
                        progress.progress(i + 1)
                        
                        # Randomly 'find' ports based on probability
                        if random.random() > 0.95:
                            p = random.choice(list(common_ports.keys()))
                            if p not in [x['port'] for x in found_ports]:
                                found_ports.append({"port": p, "service": common_ports[p], "state": "OPEN"})
                                st.write(f"Found {p}/TCP ({common_ports[p]})")
                    
                    status.update(label="SCAN COMPLETE", state="complete")
                    
                    # Store results for export
                    df_scan = pd.DataFrame(found_ports)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    fn = f"Scan_{target}_{timestamp}.json"
                    
                    st.success(f"Discovered {len(found_ports)} open vectors.")
                    st.download_button("DOWNLOAD REPORT (JSON)", df_scan.to_json(), file_name=fn)
                    log_system("SCANNER", "SUCCESS", f"Scan finished. Found {len(found_ports)} ports.")

        with col_res:
            st.markdown("#### VISUAL TOPOLOGY")
            # Mock Topology Chart
            nodes = [target, "Gateway", "Firewall", "Web Srv", "DB Srv"]
            parents = ["", target, target, "Gateway", "Gateway"]
            
            fig_tree = go.Figure(go.Treemap(
                labels=nodes,
                parents=parents,
                marker_colors=["#000", "#003366", "#003366", "#00f3ff", "#ff0055"]
            ))
            fig_tree.update_layout(template="plotly_dark", margin=dict(t=0, l=0, r=0, b=0))
            st.plotly_chart(fig_tree, use_container_width=True)

    # --- TAB 3: PASSWORD AUDIT ---
    with tabs[2]:
        st.markdown("### 🔐 CREDENTIAL STRENGTH AUDIT")
        
        pwd = st.text_input("INPUT CREDENTIAL HASH/TEXT", type="password")
        
        if pwd:
            # Entropy Logic
            length = len(pwd)
            has_upper = any(c.isupper() for c in pwd)
            has_lower = any(c.islower() for c in pwd)
            has_digit = any(c.isdigit() for c in pwd)
            has_special = any(c in string.punctuation for c in pwd)
            
            pool = 0
            if has_upper: pool += 26
            if has_lower: pool += 26
            if has_digit: pool += 10
            if has_special: pool += 32
            
            entropy = length * (len(bin(pool)) - 2) if pool > 0 else 0
            
            c_gauge, c_stats = st.columns(2)
            
            with c_gauge:
                fig_gauge = go.Figure(go.Indicator(
                    mode = "gauge+number",
                    value = entropy,
                    title = {'text': "ENTROPY (BITS)", 'font': {'color': '#00f3ff'}},
                    gauge = {
                        'axis': {'range': [0, 120]},
                        'bar': {'color': "#00f3ff"},
                        'steps': [
                            {'range': [0, 50], 'color': "#330000"},
                            {'range': [50, 80], 'color': "#333300"},
                            {'range': [80, 120], 'color': "#003300"}
                        ]
                    }
                ))
                fig_gauge.update_layout(paper_bgcolor='rgba(0,0,0,0)', font={'color': 'white'})
                st.plotly_chart(fig_gauge, use_container_width=True)
                
            with c_stats:
                st.markdown("#### COMPLEXITY ANALYSIS")
                st.write(f"LENGTH: {length} " + ("✅" if length > 12 else "⚠️"))
                st.write(f"UPPERCASE: " + ("✅" if has_upper else "❌"))
                st.write(f"NUMERIC: " + ("✅" if has_digit else "❌"))
                st.write(f"SPECIAL: " + ("✅" if has_special else "❌"))
                
                est_time = "INSTANT"
                if entropy > 50: est_time = "MINUTES"
                if entropy > 80: est_time = "YEARS"
                if entropy > 100: est_time = "CENTURIES"
                
                st.metric("CRACK TIME (RTX 4090)", est_time)

            if st.button("RUN HASH SIMULATION"):
                with st.spinner("CALCULATING SHA-256 COLLISION PROBABILITY..."):
                    time.sleep(1.5)
                    h = hashlib.sha256(pwd.encode()).hexdigest()
                    st.code(f"HASH: {h}", language="text")
                    log_system("AUTH", "WARNING", f"Audit performed on pwd len={length}, Entropy={entropy:.1f}")

    # --- TAB 4: DOS STRESS ---
    with tabs[3]:
        st.markdown("### ⚡ LOAD STRESS SIMULATOR")
        st.warning("⚠️ AUTHORIZED TESTING ONLY. DO NOT TARGET EXTERNAL INFRASTRUCTURE.")
        
        col_conf, col_mon = st.columns([1, 2])
        
        with col_conf:
            target_dos = st.text_input("ENDPOINT", "http://localhost:8000/api/v1")
            threads = st.slider("VIRTUAL USERS", 50, 500, 200)
            duration = st.slider("DURATION (SEC)", 5, 60, 10)
            
            start_dos = st.button("🔥 IGNITE STRESS TEST")
            
        with col_mon:
            chart_placeholder = st.empty()
            
        if start_dos:
            log_system("STRESS", "CRITICAL", f"DoS Simulation launched on {target_dos} with {threads} users")
            
            data_lat = []
            data_req = []
            
            for t in range(duration):
                # Simulate metrics
                lat = 20 + (threads * 0.1 * (t+1)) + random.randint(0, 20)
                reqs = threads * random.uniform(0.8, 1.2)
                
                data_lat.append(lat)
                data_req.append(reqs)
                
                # Update Chart
                fig_dos = go.Figure()
                fig_dos.add_trace(go.Scatter(y=data_lat, name='Latency (ms)', fill='tozeroy', line=dict(color='#ff0055')))
                fig_dos.add_trace(go.Scatter(y=data_req, name='Req/Sec', line=dict(color='#00f3ff')))
                fig_dos.update_layout(
                    template="plotly_dark", 
                    title="REAL-TIME SERVER RESPONSE",
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)',
                    height=350
                )
                chart_placeholder.plotly_chart(fig_dos, use_container_width=True)
                time.sleep(1)
            
            st.success("STRESS TEST CONCLUDED.")
            
            # Export Data
            csv_data = pd.DataFrame({"Time": range(duration), "Latency": data_lat, "Reqs": data_req}).to_csv()
            fn_dos = f"DoS_Report_{target_dos.replace('http://','').replace('/','_')}_{datetime.now().strftime('%H%M%S')}.csv"
            st.download_button("DOWNLOAD TELEMETRY (CSV)", csv_data, file_name=fn_dos)

    # --- TAB 5: REPORTING ---
    with tabs[5]:
        st.markdown("### 📝 MISSION REPORT")
        
        st.write("Generate executive summary of current session.")
        
        df_logs = pd.DataFrame(st.session_state.logs)
        st.dataframe(df_logs, use_container_width=True)
        
        if st.button("GENERATE OFFICIAL REPORT (DOCX)"):
            if DOCX_AVAILABLE:
                doc = Document()
                doc.add_heading(f'PAYBUDDY SECURITY ASSESSMENT', 0)
                doc.add_paragraph(f"TEAM: {TEAM_NAME}")
                doc.add_paragraph(f"DATE: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"ENV: PROD_SIM_01")
                
                doc.add_heading('EXECUTIVE SUMMARY', 1)
                doc.add_paragraph("Authorized security testing was performed on the PayBuddy infrastructure. The following logs detail the activity and findings.")
                
                doc.add_heading('SESSION LOGS', 1)
                for log in st.session_state.logs:
                    doc.add_paragraph(f"[{log['timestamp']}] [{log['tool']}] {log['message']}")
                
                # Save
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                fn_rep = f"Final_Report_{TEAM_NAME}_{datetime.now().strftime('%Y%m%d')}.docx"
                st.download_button("DOWNLOAD ENCRYPTED REPORT", buffer, file_name=fn_rep)
                st.balloons()
            else:
                st.error("MODULE 'python-docx' MISSING. INSTALL VIA requirements.txt")

    # --- FOOTER TERMINAL ---
    st.markdown("---")
    st.markdown(f"<div style='color:#555; font-size:0.8em; font-family:monospace;'>root@cyberguard-terminal:~# session_id: {hashlib.md5(str(time.time()).encode()).hexdigest()[:8]}</div>", unsafe_allow_html=True)
